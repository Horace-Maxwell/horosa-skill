from __future__ import annotations

import hashlib
import html
import json
import os
import re
import tempfile
import zipfile
from pathlib import Path
from typing import Any


HUMAN_RENDER_HIDDEN_SECTION_IDS = {
    "report_metadata",
    "report_quality",
    "delivery_checklist",
    "coverage_contract",
    "section_coverage_matrix",
    "targeted_analysis_contract",
    "question_analysis",
    "input_overview",
    "technique_summary",
    "ai_interpretation",
    "recommendations_limitations",
    "xingque_export_text",
    "provenance",
}


def render_report(document: dict[str, Any], *, output_path: Path, format_name: str) -> dict[str, Any]:
    normalized = format_name.lower()
    if normalized not in {"json", "docx", "pdf", "markdown"}:
        raise ValueError("format must be one of: json, docx, pdf, markdown")
    output_path.parent.mkdir(parents=True, exist_ok=True)

    # Render to a temp sibling first, then atomically os.replace() into place. A mid-render failure
    # (python-docx / reportlab raising, disk error) thus can never leave a truncated/corrupt artifact
    # at output_path — the destination is only touched once a complete file exists.
    fd, tmp_name = tempfile.mkstemp(dir=str(output_path.parent), prefix=f".{output_path.name}.", suffix=".tmp")
    os.close(fd)
    tmp_path = Path(tmp_name)
    try:
        if normalized == "json":
            tmp_path.write_text(json.dumps(document, ensure_ascii=False, indent=2), encoding="utf-8")
        elif normalized == "markdown":
            # markdown 只服务**技法依据报告**（确定性元数据）。咨询报告不走这条：它的正文样式规矩
            # 由 DOCX/PDF 渲染器承载，且 references/reports.md 明令正文不带机器元数据。
            from horosa_skill.reports.technique_card import render_technique_report_markdown

            tmp_path.write_text(render_technique_report_markdown(document), encoding="utf-8")
        elif normalized == "docx":
            _render_docx(document, tmp_path)
        else:  # pdf
            _render_pdf(document, tmp_path)
        payload = tmp_path.read_bytes()
        os.replace(str(tmp_path), str(output_path))
    except BaseException as exc:
        tmp_path.unlink(missing_ok=True)
        if normalized in {"docx", "pdf"}:
            # 渲染失败降级：自动改存 TXT（纯文本全文），绝不静默丢内容——结果里如实标注降级。
            degraded = _render_txt_degraded(document, output_path)
            if degraded is not None:
                degraded["degraded_from"] = normalized
                degraded["degrade_reason"] = f"{type(exc).__name__}: {exc}"
                return degraded
        raise
    return {
        "path": str(output_path),
        "format": normalized,
        "file_size": len(payload),
        "sha256": hashlib.sha256(payload).hexdigest(),
    }


def _render_txt_degraded(document: dict[str, Any], intended_path: Path) -> dict[str, Any] | None:
    # DOCX/PDF 渲染异常时的兜底：把报告全文（plain_text 或 JSON）写成 UTF-8 TXT 落在同名 .txt。
    try:
        txt_path = intended_path.with_suffix(".txt")
        text = document.get("plain_text")
        if not isinstance(text, str) or not text.strip():
            text = json.dumps(document, ensure_ascii=False, indent=2)
        txt_path.write_text(text, encoding="utf-8")
        payload = txt_path.read_bytes()
        return {
            "path": str(txt_path),
            "format": "txt",
            "file_size": len(payload),
            "sha256": hashlib.sha256(payload).hexdigest(),
        }
    except OSError:
        return None


_MD_TABLE_SEP_RE = re.compile(r"^\s*\|?\s*:?-{2,}:?\s*(\|\s*:?-{2,}:?\s*)+\|?\s*$")


def _split_markdown_blocks(text: str) -> list[tuple[str, Any]]:
    """把正文拆成 ('text', str) 与 ('table', {'header': [...], 'rows': [[...]]}) 块序列。

    识别标准 Markdown 管道表：表头行 `| a | b |` + 分隔行 `|---|---|` + 数据行。
    非表格内容原样归入 text 块（含空行分段语义交由调用方 _paragraphs 处理）。
    """
    lines = (text or "").splitlines()
    blocks: list[tuple[str, Any]] = []
    buffer: list[str] = []

    def flush_text() -> None:
        if buffer:
            blocks.append(("text", "\n".join(buffer)))
            buffer.clear()

    def parse_row(line: str) -> list[str]:
        row = line.strip()
        if row.startswith("|"):
            row = row[1:]
        if row.endswith("|"):
            row = row[:-1]
        return [cell.strip() for cell in row.split("|")]

    i = 0
    while i < len(lines):
        line = lines[i]
        is_pipe_row = "|" in line and line.strip().startswith("|")
        if is_pipe_row and i + 1 < len(lines) and _MD_TABLE_SEP_RE.match(lines[i + 1] or ""):
            header = parse_row(line)
            rows: list[list[str]] = []
            j = i + 2
            while j < len(lines) and "|" in lines[j] and lines[j].strip().startswith("|"):
                cells = parse_row(lines[j])
                # 列数对齐表头（缺补空、多截断），保证表格矩形。
                if len(cells) < len(header):
                    cells = cells + [""] * (len(header) - len(cells))
                rows.append(cells[: len(header)])
                j += 1
            flush_text()
            blocks.append(("table", {"header": header, "rows": rows}))
            i = j
            continue
        buffer.append(line)
        i += 1
    flush_text()
    return blocks


def _render_docx(document: dict[str, Any], output_path: Path) -> None:
    try:
        from docx import Document
        from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.shared import Inches, Pt, RGBColor
    except ModuleNotFoundError:
        _render_docx_fallback(document, output_path)
        return

    doc = Document()
    font_name = _docx_font_name()
    for section in doc.sections:
        section.top_margin = Inches(0.65)
        section.bottom_margin = Inches(0.62)
        section.left_margin = Inches(0.72)
        section.right_margin = Inches(0.72)
    _docx_apply_base_styles(doc, font_name, qn, Pt, RGBColor)
    docx_content_width = Inches(6.25)

    def set_cell_shading(cell: Any, fill: str) -> None:
        tc_pr = cell._tc.get_or_add_tcPr()
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), fill)
        tc_pr.append(shading)

    def set_cell_border(cell: Any, color: str = "D8DED9", size: str = "8") -> None:
        tc_pr = cell._tc.get_or_add_tcPr()
        borders = tc_pr.first_child_found_in("w:tcBorders")
        if borders is None:
            borders = OxmlElement("w:tcBorders")
            tc_pr.append(borders)
        for edge in ["top", "left", "bottom", "right"]:
            tag = f"w:{edge}"
            element = borders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                borders.append(element)
            element.set(qn("w:val"), "single")
            element.set(qn("w:sz"), size)
            element.set(qn("w:color"), color)

    def set_table_full_width(table: Any) -> None:
        tbl_pr = table._tbl.tblPr
        tbl_w = tbl_pr.find(qn("w:tblW"))
        if tbl_w is None:
            tbl_w = OxmlElement("w:tblW")
            tbl_pr.append(tbl_w)
        tbl_w.set(qn("w:type"), "pct")
        tbl_w.set(qn("w:w"), "5000")
        layout = tbl_pr.find(qn("w:tblLayout"))
        if layout is None:
            layout = OxmlElement("w:tblLayout")
            tbl_pr.append(layout)
        layout.set(qn("w:type"), "fixed")

    def set_left_border(cell: Any, color: str = "0F766E", size: str = "28") -> None:
        tc_pr = cell._tc.get_or_add_tcPr()
        borders = tc_pr.first_child_found_in("w:tcBorders")
        if borders is None:
            borders = OxmlElement("w:tcBorders")
            tc_pr.append(borders)
        left = borders.find(qn("w:left"))
        if left is None:
            left = OxmlElement("w:left")
            borders.append(left)
        left.set(qn("w:val"), "single")
        left.set(qn("w:sz"), size)
        left.set(qn("w:color"), color)

    def style_run(run: Any, *, size: int | float | None = None, bold: bool = False, color: str | None = None) -> None:
        run.font.name = font_name
        run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
        run.bold = bold
        if size is not None:
            run.font.size = Pt(size)
        if color:
            run.font.color.rgb = RGBColor.from_string(color)

    def add_data_table(header: list[str], rows: list[list[str]]) -> None:
        # Markdown 管道表 → Word 真表格：表头底纹加粗 + 跨页重复（w:tblHeader）+ 全宽 + 细边框。
        cols = max(1, len(header))
        table = doc.add_table(rows=1, cols=cols)
        table.autofit = False
        set_table_full_width(table)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        col_width = Inches(6.25 / cols)
        for column in table.columns:
            column.width = col_width
        header_row = table.rows[0]
        # 跨页重复表头
        tr_pr = header_row._tr.get_or_add_trPr()
        tbl_header = OxmlElement("w:tblHeader")
        tbl_header.set(qn("w:val"), "true")
        tr_pr.append(tbl_header)
        for idx, cell in enumerate(header_row.cells):
            cell.width = col_width
            set_cell_shading(cell, "EEF6F2")
            set_cell_border(cell)
            p = cell.paragraphs[0]
            run = p.add_run(header[idx] if idx < len(header) else "")
            style_run(run, size=9.6, bold=True, color="0F766E")
        for data_row in rows:
            row = table.add_row()
            for idx, cell in enumerate(row.cells):
                cell.width = col_width
                set_cell_border(cell)
                p = cell.paragraphs[0]
                run = p.add_run(data_row[idx] if idx < len(data_row) else "")
                style_run(run, size=9.6, color="14211F")
        doc.add_paragraph()

    def add_text(text: str, *, style: str = "Normal") -> None:
        # 正文按块渲染：Markdown 管道表成真表格，其余成段落。
        for kind, payload in _split_markdown_blocks(text):
            if kind == "table":
                add_data_table(payload["header"], payload["rows"])
                continue
            for paragraph in _paragraphs(payload):
                p = doc.add_paragraph(style=style)
                run = p.add_run(paragraph)
                style_run(run, size=10.2, color="14211F")

    def set_outline_level(paragraph: Any, level: int) -> None:
        # 给段落挂大纲级别（0=一级）：Word 导航窗格/自动目录识别 outlineLvl，无需换 Heading 视觉样式。
        p_pr = paragraph._p.get_or_add_pPr()
        outline = OxmlElement("w:outlineLvl")
        outline.set(qn("w:val"), str(level))
        p_pr.append(outline)

    def add_heading_block(text: str, *, level: int = 1) -> None:
        table = doc.add_table(rows=1, cols=1)
        table.autofit = False
        set_table_full_width(table)
        table.columns[0].width = docx_content_width
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = table.cell(0, 0)
        cell.width = docx_content_width
        set_cell_shading(cell, "EEF6F2")
        set_cell_border(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)
        set_outline_level(p, 0 if level == 1 else 1)
        run = p.add_run(text)
        style_run(run, size=16 if level == 1 else 12.5, bold=True, color="14211F" if level == 1 else "0F766E")
        doc.add_paragraph()

    def add_card(title_text: str, body_text: str, *, accent: str = "0F766E") -> None:
        table = doc.add_table(rows=1, cols=1)
        table.autofit = False
        set_table_full_width(table)
        table.columns[0].width = docx_content_width
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = table.cell(0, 0)
        cell.width = docx_content_width
        set_cell_shading(cell, "FFFFFF")
        set_cell_border(cell, color="DFE7E1")
        set_left_border(cell, color=accent, size="28")
        p = cell.paragraphs[0]
        title_run = p.add_run(title_text)
        style_run(title_run, size=12.5, bold=True, color=accent)
        for kind, payload in _split_markdown_blocks(body_text):
            if kind == "table":
                # 卡片内表格：渲染在卡片下方为独立真表格（表格嵌单元格在窄卡内可读性差）。
                pass_through_tables.append(payload)
                continue
            for paragraph in _paragraphs(payload):
                p = cell.add_paragraph()
                run = p.add_run(paragraph)
                style_run(run, size=10.2, color="14211F")
        doc.add_paragraph()
        while pass_through_tables:
            payload = pass_through_tables.pop(0)
            add_data_table(payload["header"], payload["rows"])

    pass_through_tables: list[dict[str, Any]] = []

    ai_report = document.get("ai_report") if isinstance(document.get("ai_report"), dict) else {}
    display_question = _display_question(document, ai_report)
    human_executive_summary = _human_ai_value(ai_report, "executive_summary")
    human_consultation_basis = _human_ai_value(ai_report, "consultation_basis")
    human_reading_steps = _human_ai_list(ai_report, "reading_steps")
    human_analysis_sections = _human_ai_list(ai_report, "analysis_sections")
    human_recommendations = _human_ai_list(ai_report, "recommendations")
    human_limitations = _human_ai_list(ai_report, "limitations")
    cover = doc.add_table(rows=3, cols=1)
    cover.autofit = False
    set_table_full_width(cover)
    cover.columns[0].width = docx_content_width
    cover.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row in cover.rows:
        cell = row.cells[0]
        cell.width = docx_content_width
        set_cell_shading(cell, "11231F")
        set_cell_border(cell, color="24443E", size="10")
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = cover.cell(0, 0).paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("专业咨询报告")
    style_run(run, size=9.5, color="DBE8E3")
    p = cover.cell(1, 0).paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(str(document.get("title") or "结构化咨询报告"))
    style_run(run, size=27, bold=True, color="FFFFFF")
    p = cover.cell(2, 0).paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta = "结构化解读 · 行动建议 · 风险提示"
    run = p.add_run(meta)
    style_run(run, size=9.5, color="DBE8E3")
    doc.add_paragraph()

    # 目录：TOC 域（取大纲级别 1-2），Word 打开时经 updateFields 自动生成、可点击跳转。
    toc_head = doc.add_paragraph()
    toc_run = toc_head.add_run("目录")
    style_run(toc_run, size=13, bold=True, color="0F766E")
    toc_paragraph = doc.add_paragraph()
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), 'TOC \\o "1-2" \\h \\z \\u')
    placeholder = OxmlElement("w:r")
    placeholder_text = OxmlElement("w:t")
    placeholder_text.text = "（打开文档后目录将自动更新）"
    placeholder.append(placeholder_text)
    fld.append(placeholder)
    toc_paragraph._p.append(fld)
    settings_element = doc.settings.element
    update_fields = settings_element.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings_element.append(update_fields)
    update_fields.set(qn("w:val"), "true")
    doc.add_paragraph()

    add_heading_block("解读目标")
    add_text(display_question)

    if human_consultation_basis:
        add_card("起盘依据", str(human_consultation_basis or ""), accent="0F766E")
    if human_reading_steps:
        add_card("解盘步骤", _steps_text(human_reading_steps), accent="2563EB")
    if ai_report.get("analysis_focus") or ai_report.get("direct_answer"):
        add_card(
            "核心结论",
            str(ai_report.get("direct_answer") or ai_report.get("analysis_focus") or ""),
        )
    if human_executive_summary:
        add_heading_block("总览摘要")
        add_text(str(human_executive_summary or ""))
    if ai_report.get("answer_text") and str(ai_report.get("answer_text")) != str(human_executive_summary or ""):
        add_heading_block("完整解盘正文")
        add_text(str(ai_report.get("answer_text") or ""))
    if human_analysis_sections:
        add_heading_block("正文解读")
        for section in human_analysis_sections:
            if isinstance(section, dict):
                add_heading_block(str(section.get("title") or "分析"), level=2)
                add_text(str(section.get("body") or section.get("content") or ""))
            else:
                add_text(str(section))
    if human_recommendations:
        add_card("建议", "\n".join(f"- {item}" for item in human_recommendations), accent="B7791F")
    if human_limitations:
        add_heading_block("限制与注意事项")
        add_text("\n".join(f"- {item}" for item in human_limitations))

    for section in _human_sections(document):
        if not isinstance(section, dict):
            continue
        add_heading_block(str(section.get("title") or "章节"))
        add_text(str(section.get("body") or ""))

    # 文档元数据：文件属性可检索/归档（标题/作者/主题/关键词/创建时间）。
    # OOXML core property 单项硬限 255 字符；search_index 里可能混入整段长文本，先滤短词再截断。
    source_meta = document.get("source") if isinstance(document.get("source"), dict) else {}
    props = doc.core_properties
    props.title = _core_prop_text(document.get("title") or "结构化咨询报告")
    props.author = "Horosa Skill"
    props.subject = _core_prop_text(source_meta.get("technique_label") or source_meta.get("tool_name") or "")
    keywords = document.get("search_index") if isinstance(document.get("search_index"), dict) else {}
    keyword_list = keywords.get("keywords") if isinstance(keywords.get("keywords"), list) else []
    short_keywords = [str(k).strip() for k in keyword_list if str(k).strip() and len(str(k).strip()) <= 24]
    props.keywords = _core_prop_text(", ".join(short_keywords[:12]))
    props.comments = "Generated by Horosa Skill report renderer"

    # 页脚：报告标题 + 「第 X 页 / 共 Y 页」域（与 PDF 页脚对齐）。
    footer = doc.sections[0].footer
    footer_p = footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = footer_p.add_run(f"{props.title}　·　第 ")
    style_run(title_run, size=8, color="66706D")
    for instr in ("PAGE", "NUMPAGES"):
        fld_page = OxmlElement("w:fldSimple")
        fld_page.set(qn("w:instr"), instr)
        run_el = OxmlElement("w:r")
        text_el = OxmlElement("w:t")
        text_el.text = "1"
        run_el.append(text_el)
        fld_page.append(run_el)
        footer_p._p.append(fld_page)
        if instr == "PAGE":
            mid_run = footer_p.add_run(" 页 / 共 ")
            style_run(mid_run, size=8, color="66706D")
    tail_run = footer_p.add_run(" 页")
    style_run(tail_run, size=8, color="66706D")

    doc.save(str(output_path))


def _render_pdf(document: dict[str, Any], output_path: Path) -> None:
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib import colors
        from reportlab.lib.enums import TA_CENTER
        from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
        from reportlab.lib.units import mm
        from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.cidfonts import UnicodeCIDFont
        from reportlab.pdfbase.ttfonts import TTFont
    except ModuleNotFoundError:
        _render_pdf_fallback(document, output_path)
        return

    font_name, bold_font_name = _register_pdf_fonts(pdfmetrics, UnicodeCIDFont, TTFont)

    styles = getSampleStyleSheet()
    page_width, page_height = A4
    left_margin = right_margin = 22 * mm
    top_margin = bottom_margin = 18 * mm
    content_width = page_width - left_margin - right_margin
    ink = colors.HexColor("#14211f")
    muted = colors.HexColor("#66706d")
    green = colors.HexColor("#0f766e")
    amber = colors.HexColor("#b7791f")
    paper = colors.HexColor("#fbfaf5")
    pale = colors.HexColor("#eef6f2")
    line = colors.HexColor("#d8ded9")

    base = ParagraphStyle(
        "HorosaBody",
        parent=styles["BodyText"],
        fontName=font_name,
        fontSize=9.6,
        leading=16,
        textColor=ink,
        spaceAfter=4,
    )
    small = ParagraphStyle("HorosaSmall", parent=base, fontSize=8.2, leading=12, textColor=muted)
    cover_title = ParagraphStyle(
        "HorosaCoverTitle",
        parent=styles["Title"],
        fontName=bold_font_name,
        fontSize=27,
        leading=34,
        textColor=colors.white,
        alignment=TA_CENTER,
    )
    cover_meta = ParagraphStyle(
        "HorosaCoverMeta",
        parent=base,
        fontName=font_name,
        fontSize=10.5,
        leading=16,
        textColor=colors.HexColor("#dbe8e3"),
        alignment=TA_CENTER,
    )
    heading = ParagraphStyle(
        "HorosaHeading",
        parent=styles["Heading1"],
        fontName=bold_font_name,
        fontSize=15,
        leading=21,
        textColor=ink,
        spaceBefore=14,
        spaceAfter=8,
    )
    subheading = ParagraphStyle(
        "HorosaSubHeading",
        parent=styles["Heading2"],
        fontName=bold_font_name,
        fontSize=11.5,
        leading=17,
        textColor=green,
        spaceBefore=8,
        spaceAfter=5,
    )
    story: list[Any] = []

    def add_heading(text: str, style: ParagraphStyle = heading) -> None:
        block = Table(
            [[Paragraph(_esc(text), style)]],
            colWidths=[content_width],
            style=TableStyle(
                [
                    ("BACKGROUND", (0, 0), (-1, -1), pale),
                    ("BOX", (0, 0), (-1, -1), 0.55, line),
                    ("LEFTPADDING", (0, 0), (-1, -1), 9),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 9),
                    ("TOPPADDING", (0, 0), (-1, -1), 6),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
                ]
            ),
        )
        story.append(Spacer(1, 8))
        story.append(block)
        story.append(Spacer(1, 5))

    table_cell = ParagraphStyle("HorosaTableCell", parent=base, fontSize=8.6, leading=12, spaceAfter=0)
    table_head = ParagraphStyle("HorosaTableHead", parent=table_cell, fontName=bold_font_name, textColor=green)

    def add_data_table(header: list[str], rows: list[list[str]]) -> None:
        # Markdown 管道表 → PDF 真表格：表头底纹 + repeatRows 跨页重复表头 + 细网格。
        cols = max(1, len(header))
        data = [[Paragraph(_esc(h), table_head) for h in header]]
        for row in rows:
            data.append([Paragraph(_esc(row[idx] if idx < len(row) else ""), table_cell) for idx in range(cols)])
        story.append(
            Table(
                data,
                colWidths=[content_width / cols] * cols,
                repeatRows=1,
                style=TableStyle(
                    [
                        ("BACKGROUND", (0, 0), (-1, 0), pale),
                        ("GRID", (0, 0), (-1, -1), 0.4, line),
                        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                        ("LEFTPADDING", (0, 0), (-1, -1), 5),
                        ("RIGHTPADDING", (0, 0), (-1, -1), 5),
                        ("TOPPADDING", (0, 0), (-1, -1), 3),
                        ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
                    ]
                ),
            )
        )
        story.append(Spacer(1, 6))

    def add_body(text: str) -> None:
        for kind, payload in _split_markdown_blocks(text):
            if kind == "table":
                add_data_table(payload["header"], payload["rows"])
                continue
            for paragraph in _paragraphs(payload):
                story.append(Paragraph(_esc(paragraph), base))
                story.append(Spacer(1, 4))

    def add_note_card(title_text: str, body_text: str, *, accent: Any = green) -> None:
        rows = [
            [Paragraph(_esc(title_text), ParagraphStyle("CardTitle", parent=subheading, textColor=accent))],
            *[[Paragraph(_esc(paragraph), base)] for paragraph in _paragraphs(body_text)],
        ]
        story.append(
            Table(
                rows,
                colWidths=[content_width],
                style=TableStyle(
                    [
                        ("BACKGROUND", (0, 0), (-1, -1), colors.white),
                        ("BOX", (0, 0), (-1, -1), 0.6, colors.HexColor("#dfe7e1")),
                        ("LINEBEFORE", (0, 0), (0, -1), 3, accent),
                        ("LEFTPADDING", (0, 0), (-1, -1), 11),
                        ("RIGHTPADDING", (0, 0), (-1, -1), 11),
                        ("TOPPADDING", (0, 0), (-1, -1), 7),
                        ("BOTTOMPADDING", (0, 0), (-1, -1), 7),
                    ]
                ),
            )
        )
        story.append(Spacer(1, 9))

    ai_report = document.get("ai_report") if isinstance(document.get("ai_report"), dict) else {}
    display_question = _display_question(document, ai_report)
    human_executive_summary = _human_ai_value(ai_report, "executive_summary")
    human_consultation_basis = _human_ai_value(ai_report, "consultation_basis")
    human_reading_steps = _human_ai_list(ai_report, "reading_steps")
    human_analysis_sections = _human_ai_list(ai_report, "analysis_sections")
    human_recommendations = _human_ai_list(ai_report, "recommendations")
    human_limitations = _human_ai_list(ai_report, "limitations")
    metadata = ["结构化解读 · 行动建议 · 风险提示"]
    cover = Table(
        [
            [Paragraph("专业咨询报告", cover_meta)],
            [Paragraph(_esc(str(document.get("title") or "结构化咨询报告")), cover_title)],
            [Paragraph(_esc(" / ".join(metadata)), cover_meta)],
        ],
        colWidths=[content_width],
        style=TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#11231f")),
                ("BOX", (0, 0), (-1, -1), 0.8, colors.HexColor("#24443e")),
                ("LEFTPADDING", (0, 0), (-1, -1), 20),
                ("RIGHTPADDING", (0, 0), (-1, -1), 20),
                ("TOPPADDING", (0, 0), (0, 0), 20),
                ("BOTTOMPADDING", (0, 0), (0, 0), 8),
                ("TOPPADDING", (0, 1), (0, 1), 10),
                ("BOTTOMPADDING", (0, 1), (0, 1), 12),
                ("TOPPADDING", (0, 2), (0, 2), 8),
                ("BOTTOMPADDING", (0, 2), (0, 2), 20),
            ]
        ),
    )
    story.append(cover)
    story.append(Spacer(1, 14))

    add_heading("解读目标")
    add_body(display_question)
    if human_consultation_basis:
        add_note_card("起盘依据", str(human_consultation_basis or ""), accent=green)
    if human_reading_steps:
        add_note_card("解盘步骤", _steps_text(human_reading_steps), accent=colors.HexColor("#2563eb"))
    if ai_report.get("analysis_focus") or ai_report.get("direct_answer"):
        add_note_card(
            "核心结论",
            str(ai_report.get("direct_answer") or ai_report.get("analysis_focus") or ""),
            accent=green,
        )
    if human_executive_summary:
        add_heading("总览摘要")
        add_body(str(human_executive_summary or ""))
    if ai_report.get("answer_text") and str(ai_report.get("answer_text")) != str(human_executive_summary or ""):
        add_heading("完整解盘正文")
        add_body(str(ai_report.get("answer_text") or ""))
    if human_analysis_sections:
        add_heading("正文解读")
        for section in human_analysis_sections:
            if isinstance(section, dict):
                add_heading(str(section.get("title") or "分析"), subheading)
                body_parts = []
                body_parts.append(str(section.get("body") or section.get("content") or ""))
                add_body("\n".join(body_parts))
            else:
                add_body(str(section))
    if human_recommendations:
        add_note_card("建议", "\n".join(f"- {item}" for item in human_recommendations), accent=amber)
    if human_limitations:
        add_heading("限制与注意事项")
        add_body("\n".join(f"- {item}" for item in human_limitations))

    for section in _human_sections(document):
        if not isinstance(section, dict):
            continue
        add_heading(str(section.get("title") or "章节"))
        add_body(str(section.get("body") or ""))

    def on_page(canvas: Any, doc: Any) -> None:
        canvas.saveState()
        canvas.setFillColor(paper)
        canvas.rect(0, 0, page_width, page_height, fill=1, stroke=0)
        canvas.setFillColor(colors.HexColor("#e6eee9"))
        canvas.rect(0, page_height - 8, page_width, 8, fill=1, stroke=0)
        canvas.setFont(font_name, 7.5)
        canvas.setFillColor(muted)
        canvas.drawString(left_margin, 10 * mm, str(document.get("title") or "咨询报告"))
        canvas.drawRightString(page_width - right_margin, 10 * mm, f"第 {doc.page} 页")
        canvas.restoreState()

    SimpleDocTemplate(
        str(output_path),
        pagesize=A4,
        rightMargin=right_margin,
        leftMargin=left_margin,
        topMargin=top_margin,
        bottomMargin=bottom_margin,
    ).build(story, onFirstPage=on_page, onLaterPages=on_page)


def _docx_font_name() -> str:
    # 中文字体候选（按平台常见优先级）：Win 微软雅黑 → Mac 苹方/宋体 → Arial Unicode → Noto CJK。
    # 返回的字体名同时用于 w:eastAsia，确保 CJK 字形真实存在，异机打开不豆腐。
    if Path("C:/Windows/Fonts/msyh.ttc").exists() or Path("C:/Windows/Fonts/msyh.ttf").exists():
        return "Microsoft YaHei"
    if Path("/System/Library/Fonts/PingFang.ttc").exists():
        return "PingFang SC"
    if Path("/System/Library/Fonts/Supplemental/Songti.ttc").exists():
        return "Songti SC"
    if Path("/System/Library/Fonts/Supplemental/Arial Unicode.ttf").exists():
        return "Arial Unicode MS"
    for noto in ("/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc", "/usr/share/fonts/noto-cjk/NotoSansCJK-Regular.ttc"):
        if Path(noto).exists():
            return "Noto Sans CJK SC"
    return "Arial"


def _docx_apply_base_styles(doc: Any, font_name: str, qn: Any, Pt: Any, RGBColor: Any) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = font_name
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    normal.font.size = Pt(10.2)
    normal.font.color.rgb = RGBColor.from_string("14211F")
    for style_name, size, color in [
        ("Title", 27, "14211F"),
        ("Heading 1", 16, "14211F"),
        ("Heading 2", 12.5, "0F766E"),
        ("List Bullet", 9.5, "66706D"),
    ]:
        try:
            style = styles[style_name]
        except KeyError:
            continue
        style.font.name = font_name
        style._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)


def _docx_heading(doc: Any, text: str, level: int) -> None:
    doc.add_heading(text, level=level)


def _paragraphs(text: str) -> list[str]:
    raw = str(text or "").replace("\r\n", "\n").replace("\r", "\n")
    paragraphs = [item.strip() for item in raw.split("\n") if item.strip()]
    return paragraphs or ["无"]


def _esc(value: str) -> str:
    return html.escape(str(value)).replace("\n", "<br/>")


def _register_pdf_fonts(pdfmetrics: Any, UnicodeCIDFont: Any, TTFont: Any) -> tuple[str, str]:
    candidates = [
        "C:/Windows/Fonts/msyh.ttc",
        "C:/Windows/Fonts/msyh.ttf",
        "C:/Windows/Fonts/simhei.ttf",
        "C:/Windows/Fonts/simsun.ttc",
        "C:/Windows/Fonts/simsun.ttf",
        # macOS 中文正体候选（苹方/宋体优先于 Arial Unicode）
        "/System/Library/Fonts/PingFang.ttc",
        "/System/Library/Fonts/Supplemental/Songti.ttc",
        "/System/Library/Fonts/STHeiti Light.ttc",
        "/System/Library/Fonts/Supplemental/Arial Unicode.ttf",
        "/System/Library/Fonts/Supplemental/AppleGothic.ttf",
        "/System/Library/Fonts/SFNS.ttf",
        # Linux Noto CJK
        "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc",
        "/usr/share/fonts/noto-cjk/NotoSansCJK-Regular.ttc",
    ]
    for index, candidate in enumerate(candidates):
        path = Path(candidate)
        if not path.exists():
            continue
        font_name = f"HorosaSans{index}"
        try:
            if path.suffix.lower() == ".ttc":
                # TTC 集合字体需指定子字体索引（0=常规），否则 reportlab 拒载。
                pdfmetrics.registerFont(TTFont(font_name, str(path), subfontIndex=0))
            else:
                pdfmetrics.registerFont(TTFont(font_name, str(path)))
            return font_name, font_name
        except Exception:
            continue
    fallback = "STSong-Light"
    try:
        pdfmetrics.registerFont(UnicodeCIDFont(fallback))
        return fallback, fallback
    except Exception:
        return "Helvetica", "Helvetica"


def _core_prop_text(value: Any, limit: int = 255) -> str:
    """OOXML core property 文本硬限 255 字符，超限写入会直接抛错，统一在此截断。"""
    text = str(value or "").strip()
    return text if len(text) <= limit else text[: limit - 1] + "…"


def _human_sections(document: dict[str, Any]) -> list[dict[str, Any]]:
    sections = document.get("sections") if isinstance(document.get("sections"), list) else []
    output: list[dict[str, Any]] = []
    for section in sections:
        if not isinstance(section, dict):
            continue
        if section.get("human_visible") is not True:
            continue
        if str(section.get("id") or "") in HUMAN_RENDER_HIDDEN_SECTION_IDS:
            continue
        output.append(section)
    return output


def _human_ai_value(ai_report: dict[str, Any], key: str) -> Any:
    human_key = f"human_{key}"
    if human_key in ai_report:
        return ai_report.get(human_key)
    return ai_report.get(key)


def _human_ai_list(ai_report: dict[str, Any], key: str) -> list[Any]:
    value = _human_ai_value(ai_report, key)
    return value if isinstance(value, list) else []


def _display_question(document: dict[str, Any], ai_report: dict[str, Any]) -> str:
    raw_question = str(document.get("user_question") or "").strip()
    if raw_question:
        return raw_question
    focus = str(ai_report.get("analysis_focus") or "").strip()
    return focus or "整体综合解盘"


def _steps_text(value: Any) -> str:
    if isinstance(value, list):
        return "\n".join(f"{index}. {item}" for index, item in enumerate(value, start=1) if str(item).strip())
    return str(value or "")


def _render_docx_fallback(document: dict[str, Any], output_path: Path) -> None:
    ai_report = document.get("ai_report") if isinstance(document.get("ai_report"), dict) else {}
    human_executive_summary = _human_ai_value(ai_report, "executive_summary")
    human_consultation_basis = _human_ai_value(ai_report, "consultation_basis")
    human_reading_steps = _human_ai_list(ai_report, "reading_steps")
    paragraphs: list[tuple[str, str]] = [("Title", str(document.get("title") or "结构化咨询报告"))]
    paragraphs.extend(
        [
            ("Heading1", "解读目标"),
            ("Normal", _display_question(document, ai_report)),
        ]
    )
    if human_consultation_basis:
        paragraphs.append(("Heading1", "起盘依据"))
        paragraphs.append(("Normal", str(human_consultation_basis or "")))
    if human_reading_steps:
        paragraphs.append(("Heading1", "解盘步骤"))
        paragraphs.append(("Normal", _steps_text(human_reading_steps)))
    if human_executive_summary:
        paragraphs.append(("Heading1", "总览摘要"))
        paragraphs.append(("Normal", str(human_executive_summary or "")))
    if ai_report.get("answer_text") and str(ai_report.get("answer_text")) != str(human_executive_summary or ""):
        paragraphs.append(("Heading1", "完整解盘正文"))
        paragraphs.append(("Normal", str(ai_report.get("answer_text") or "")))
    if ai_report.get("analysis_focus") or ai_report.get("direct_answer"):
        paragraphs.append(("Heading1", "核心结论"))
        if ai_report.get("analysis_focus"):
            paragraphs.append(("Normal", f"分析焦点：{ai_report.get('analysis_focus')}"))
        if ai_report.get("direct_answer"):
            paragraphs.append(("Normal", f"直接回答：{ai_report.get('direct_answer')}"))
    for section in _human_sections(document):
        if isinstance(section, dict):
            paragraphs.append(("Heading1", str(section.get("title") or "章节")))
            paragraphs.extend(("Normal", item) for item in _paragraphs(str(section.get("body") or "")))
    document_xml = "\n".join(_docx_paragraph(style, text) for style, text in paragraphs)
    with zipfile.ZipFile(output_path, "w", compression=zipfile.ZIP_DEFLATED) as archive:
        archive.writestr("[Content_Types].xml", _DOCX_CONTENT_TYPES)
        archive.writestr("_rels/.rels", _DOCX_RELS)
        archive.writestr("word/_rels/document.xml.rels", _DOCX_DOCUMENT_RELS)
        archive.writestr("word/styles.xml", _DOCX_STYLES)
        archive.writestr("word/document.xml", f"{_DOCX_DOCUMENT_PREFIX}{document_xml}{_DOCX_DOCUMENT_SUFFIX}")


def _docx_paragraph(style: str, text: str) -> str:
    style_xml = f'<w:pPr><w:pStyle w:val="{style}"/></w:pPr>' if style != "Normal" else ""
    runs = "".join(f"<w:r><w:t>{html.escape(part)}</w:t></w:r>" for part in str(text).split("\n"))
    return f"<w:p>{style_xml}{runs}</w:p>"


def _render_pdf_fallback(document: dict[str, Any], output_path: Path) -> None:
    ai_report = document.get("ai_report") if isinstance(document.get("ai_report"), dict) else {}
    human_executive_summary = _human_ai_value(ai_report, "executive_summary")
    human_consultation_basis = _human_ai_value(ai_report, "consultation_basis")
    human_reading_steps = _human_ai_list(ai_report, "reading_steps")
    lines = [
        str(document.get("title") or "结构化咨询报告"),
        f"解读目标: {_display_question(document, ai_report)}",
    ]
    if human_consultation_basis:
        lines.append(f"起盘依据: {human_consultation_basis}")
    if human_reading_steps:
        lines.append(f"解盘步骤: {_steps_text(human_reading_steps)}")
    if human_executive_summary:
        lines.append(f"总览摘要: {human_executive_summary}")
    if ai_report.get("answer_text") and str(ai_report.get("answer_text")) != str(human_executive_summary or ""):
        lines.append(f"完整解盘正文: {ai_report.get('answer_text')}")
    for section in _human_sections(document)[:8]:
        if isinstance(section, dict):
            lines.append(str(section.get("title") or "Section"))
            lines.extend(_paragraphs(str(section.get("body") or ""))[:8])
    ascii_lines = [_pdf_safe(line)[:110] for line in lines[:60]]
    text_stream = "BT /F1 10 Tf 50 790 Td 14 TL " + " T* ".join(f"({_pdf_escape(line)}) Tj" for line in ascii_lines) + " ET"
    stream_bytes = text_stream.encode("latin-1", errors="replace")
    objects = [
        b"<< /Type /Catalog /Pages 2 0 R >>",
        b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
        b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 595 842] /Resources << /Font << /F1 4 0 R >> >> /Contents 5 0 R >>",
        b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>",
        b"<< /Length " + str(len(stream_bytes)).encode("ascii") + b" >>\nstream\n" + stream_bytes + b"\nendstream",
    ]
    chunks = [b"%PDF-1.4\n"]
    offsets: list[int] = []
    for index, obj in enumerate(objects, start=1):
        offsets.append(sum(len(chunk) for chunk in chunks))
        chunks.append(f"{index} 0 obj\n".encode("ascii") + obj + b"\nendobj\n")
    xref_offset = sum(len(chunk) for chunk in chunks)
    chunks.append(f"xref\n0 {len(objects) + 1}\n0000000000 65535 f \n".encode("ascii"))
    for offset in offsets:
        chunks.append(f"{offset:010d} 00000 n \n".encode("ascii"))
    chunks.append(f"trailer << /Size {len(objects) + 1} /Root 1 0 R >>\nstartxref\n{xref_offset}\n%%EOF\n".encode("ascii"))
    output_path.write_bytes(b"".join(chunks))


def _pdf_safe(value: str) -> str:
    normalized = re.sub(r"\s+", " ", str(value)).strip()
    return normalized.encode("latin-1", errors="replace").decode("latin-1")


def _pdf_escape(value: str) -> str:
    return str(value).replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")


_DOCX_CONTENT_TYPES = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">
  <Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>
  <Default Extension="xml" ContentType="application/xml"/>
  <Override PartName="/word/document.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/>
  <Override PartName="/word/styles.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.styles+xml"/>
</Types>"""
_DOCX_RELS = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
  <Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" Target="word/document.xml"/>
</Relationships>"""
_DOCX_DOCUMENT_RELS = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships"/>"""
_DOCX_DOCUMENT_PREFIX = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:body>"""
_DOCX_DOCUMENT_SUFFIX = """<w:sectPr><w:pgSz w:w="11906" w:h="16838"/><w:pgMar w:top="1440" w:right="1440" w:bottom="1440" w:left="1440"/></w:sectPr></w:body></w:document>"""
_DOCX_STYLES = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:styles xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
  <w:style w:type="paragraph" w:styleId="Normal"><w:name w:val="Normal"/></w:style>
  <w:style w:type="paragraph" w:styleId="Title"><w:name w:val="Title"/><w:basedOn w:val="Normal"/><w:rPr><w:b/><w:sz w:val="36"/></w:rPr></w:style>
  <w:style w:type="paragraph" w:styleId="Heading1"><w:name w:val="heading 1"/><w:basedOn w:val="Normal"/><w:rPr><w:b/><w:sz w:val="28"/></w:rPr></w:style>
</w:styles>"""
